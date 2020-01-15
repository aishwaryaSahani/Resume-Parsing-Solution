import csv
import re
import spacy
import docx
nlp = spacy.load("en_core_web_sm")

# The purpose of the method is to find named entities, phrases and concepts
def namedEntityRecognition(text):
    doc = nlp(text)
#     print(doc)
    for entity in doc.ents:
        print(entity.text, entity.label_)

# the purpose of the method is to extract the first phone number from the text. Regex have been used for this purpose.
def extractMobileNumber(text):
    phone_regex = r"\b((\+\d{1,3}[.\-\) ]*)*[(]?\d{3}[.\-\) ]*\d{3}[.\-\) ]*\d{4})\b"
    match = re.findall(phone_regex, text)
    number =""
    if match:
        phoneNo = re.findall(r'\d+', str(match[0]))
        number = ('-'.join(str(elem) for elem in phoneNo))
    return number

# the purpose of the method is to extract the first email from the text. Regex have been used for this purpose.

def extractEmailId(text):
    email_regex = r"\b([\w]+[.]?[\w]+@+[a-zA-Z]+[.][a-zA-Z]+)\b"
    match = re.findall(email_regex, text)
    email = ""
    if match:
        email = (''.join(str(elem) for elem in str(match[0]).replace("\n","")))    
    return str(email)

# the purpose of the method is to extract the first person  name from the text. Named Entity Recognition is used for this purpose.
# we fetch the different entities and the entity with PERSON as attribute is returned.
def extractPersonName(text): 
    doc = nlp(text)
    name = ""
#     properNoun = False
#     for word in doc:
#         print(word.text+" "+word.pos_+" "+str(word.dep_))
#         if(word.pos_=='PROPN'):
#             properNoun = True
#         name = name + " "+(word.text)
#         if word.pos_ == "SPACE" and len(name)>0 and properNoun:
#             return str(name.strip())
#         elif word.pos_ == "SPACE":
#             name = ""
#             
    for entity in doc.ents:
        if(entity.label_ == "PERSON"):
            return entity.text.strip()

    return name


# the purpose of the method is to extract the skill set from the text. The predicted skills sections is sent to the method as input.
# I have a dataset of 50000 skills identified from linkedin which searches thru the section to match skills.

def extractSkills(text):
    
#     The list of linkedin skills  is loaded
    f = open("all_linked_skills.txt", "r")
    skills = (f.read())
    skillsList = skills.split("\n")
    profileSkills = set()
#     profileSkills2 = set()
    
#     tokens = get_tokens(text)
#     for word in tokens:
#         if(word in skillsList):
#             profileSkills.add(word)
#             
#     for word in skillsList:
#         if word in text:
#             profileSkills2.add(word)
#     print(len(profileSkills2))
#     return str(profileSkills2)
# #     
#     removing the stopwords to improve the performance
#     doc = nlp(text)
#     tokens = [token.text for token in doc if not token.is_stop]
#     plainTextAfterStopWordRemoval = (' '.join(str(elem) for elem in tokens))
#     print((plainTextAfterStopWordRemoval))
#     print(len(doc))

    for skill in skillsList:
        if ("++") in skill:
            continue
        skills_regex = r"\b"+skill+"\\b"
        match = re.findall(skills_regex, text, re.IGNORECASE)
        if len(match) >0:
            profileSkills.add(skill) 
#     print(len(profileSkills))
    return (profileSkills)


# the purpose of the method is to extract the education from the text. The predicted education sections is sent to the method as input.
# Like person name, we use named entity recognition for the purpose and identify different ORG entites in the text and return the ones which have Universoty, COllege or similar words in them.
def extractEducation(text):
#     print(text)
#     The education section consists of COllege, From and TO year and the Degree 


    doc = nlp(text) 
    
    colleges = fetchUniversity(doc)
    dates = fetchFromToDates(doc) 
    degree = fetchDegree(text)
    return str(colleges)+str(dates)+str(degree)

# the purpose of the method is to extract the degree from the text. A major list is used for this purpose to match the majors 
def fetchDegree(text):
    degreeList = []
    candDegree = []
    with open('majors-list.csv', 'r') as csvfile:
        file = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in file:
            degreeList.append(row[0])
        for degree in degreeList:
            if(degree in text.upper()):
                candDegree.append(degree)
    return candDegree

# the purpose of the method is to extract the college name from the text.
# Like person name, we use named entity recognition for the purpose and identify different ORG entites in the text and return the ones which have Universoty, COllege or similar words in them.

def fetchUniversity(doc):
    org = []
     
    for entity in doc.ents:
        if(entity.label_ == "ORG" and ("University")in str(entity.text) or ("Institute")in str(entity.text) or ("College")in str(entity.text) or ("School")in str(entity.text)):
            org.append(entity.text.strip())
    return org


# the purpose of the method is to extract the company name from the text.
# Like person name, we use named entity recognition for the purpose and identify different ORG entites in the text
def fetchOrganization(doc):
#     print(doc)
    org = []
    for entity in doc.ents:
#         print(entity.text,entity.label_)
        if(entity.label_ == "ORG"):
            org.append(entity.text.strip())
    return org


# the purpose of the method is to extract the dates from the text.
# Like person name, we use named entity recognition for the purpose and identify different DATE entites in the text
def fetchFromToDates(doc):
    
    dates = []
    for entity in doc.ents:
        if(entity.label_ == "DATE"):
            dates.append(entity.text.strip())
            
    return dates

# the purpose of the method is to extract the experience from the text. The predicted experience sections is sent to the method as input.
# Like person name, we use named entity recognition for the purpose and identify different entites in the text

def extractExperience(text):
#     print(text)
    doc = nlp(text) 
    companies = fetchOrganization(doc)
    dates = fetchFromToDates(doc)
    position = matchPositions(text)
    return str(companies)+str(dates)+str(position)


# the purpose of the method is to extract the positions from the text. A position list is used for this purpose to match the job positions 
def matchPositions(text):
    f = open("job-phrase-list.txt", "r")
    positions = (f.read())
    positionsList = positions.split("\n")
    posSet = set()
    for pos in positionsList:
        pos_regex = pos
        
        if pos.upper() in text.upper():
            posSet.add(pos)
#         match = re.findall(pos_regex, text)
#         if len(match) >0:
#             posSet.add(pos) 
#     print(len(profileSkills))
    return (posSet)


# The purpose of this method is to find the numerous sections in the text and divide the text into sections. The idemtified section headins and indices help us in finding sections. 
#We also have a list of similar labels for each secrtion which were trained in semi supervised way to help in matching the sections.

def findSection(path, category, index, sections):
    
#     csv containing the similr labels for each section. The label names are matched with section names identified from the resume to predict the sections.
    with open('Train\Labels.csv', 'r') as csvfile:
        col = -1
        file = csv.reader(csvfile, delimiter=',', quotechar='"')
        for row in file:
            if(row[0]==category.upper()):
                className = row
                break
    doc = docx.Document(path)
    fullText = []
    start=end=count = 0
    
    
#     with the indices we define the start and end points which help us to find the section
    for i in range(len(sections)):
        if(sections[i].upper() in className):
            start = index[i]+1
            if i+1 < len(sections):
                end = index[i+1]
            else:
                end = len(doc.paragraphs)
            break
    if (start == end and end!=0):
        end = len(doc.paragraphs)
    for para in range(start,end):
        fullText.append(doc.paragraphs[para].text) 
    return ('. '.join(fullText))
